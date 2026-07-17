from types import SimpleNamespace

import pytest

from app.services import ai_brief_document_service as document_module
from app.services.document_parser_service import DocumentParseError, ParsedSection


def test_pdf_confirmation_reply_lists_extracted_fields_without_saving_them():
    reply = document_module.build_brief_document_confirmation_reply(
        document_module.BriefDocumentExtraction(
            updates={
                "project_name": "春季发布会",
                "city_location": "杭州湖滨银泰",
            },
            filenames=["brief.pdf"],
        )
    )

    assert "项目名称**：春季发布会" in reply
    assert "投放城市与媒体位置**：杭州湖滨银泰" in reply
    assert "并纳入本次 Brief" in reply
    assert "未提及的内容将保持" in reply


def test_pdf_revision_reply_lists_only_corrected_fields():
    reply = document_module.build_brief_document_revision_reply(
        {"city_location": "上海南京东路"}
    )

    assert "投放城市与媒体位置**：上海南京东路" in reply
    assert "项目名称" not in reply


@pytest.mark.asyncio
async def test_extract_uploaded_pdf_brief_fields(monkeypatch, tmp_path):
    upload_dir = tmp_path / "uploads"
    pdf_dir = upload_dir / "site_photos" / "user-test"
    pdf_dir.mkdir(parents=True)
    pdf_path = pdf_dir / "brief.pdf"
    pdf_path.write_bytes(b"%PDF-test")

    monkeypatch.setattr(document_module.settings, "UPLOAD_DIR", str(upload_dir))
    monkeypatch.setattr(document_module.settings, "OSS_ENABLED", False)
    monkeypatch.setattr(document_module.settings, "AI_API_KEY", "test-key")
    monkeypatch.setattr(
        document_module,
        "parse_document",
        lambda *_: [ParsedSection(label="第1页", page=1, text="项目名称：春季发布会；投放城市：杭州湖滨银泰")],
    )

    async def _mock_completion(payload, *, timeout=None, attempts=None):
        assert "杭州湖滨银泰" in payload["messages"][1]["content"]
        return {
            "choices": [
                {
                    "message": {
                        "content": (
                            '{"project_name":"春季发布会","city_location":"杭州湖滨银泰",'
                            '"theme_concept":"","remarks":""}'
                        )
                    }
                }
            ]
        }

    monkeypatch.setattr(document_module, "post_chat_completion", _mock_completion)
    result = await document_module.extract_uploaded_brief_documents(
        [
            SimpleNamespace(
                name="brief.pdf",
                url="/uploads/site_photos/user-test/brief.pdf",
                type="application/pdf",
                object_key="",
            )
        ],
        user_id="user-test",
    )

    assert result.updates == {
        "project_name": "春季发布会",
        "city_location": "杭州湖滨银泰",
    }
    assert document_module.BRIEF_DOCUMENT_CONTEXT_MARKER in result.context
    assert "项目名称：春季发布会" in result.context


@pytest.mark.asyncio
async def test_extract_uploaded_pdf_rejects_other_user_path(monkeypatch, tmp_path):
    monkeypatch.setattr(document_module.settings, "UPLOAD_DIR", str(tmp_path))
    monkeypatch.setattr(document_module.settings, "OSS_ENABLED", False)
    monkeypatch.setattr(document_module.settings, "AI_API_KEY", "test-key")

    result = await document_module.extract_uploaded_brief_documents(
        [
            SimpleNamespace(
                name="brief.pdf",
                url="/uploads/site_photos/another-user/brief.pdf",
                type="application/pdf",
                object_key="",
            )
        ],
        user_id="user-test",
    )

    assert not result.updates
    assert result.failures == ["brief.pdf：无法访问文件"]


@pytest.mark.asyncio
async def test_extract_uploaded_docx_brief_fields(monkeypatch, tmp_path):
    from docx import Document

    upload_dir = tmp_path / "uploads"
    docx_dir = upload_dir / "site_photos" / "user-test"
    docx_dir.mkdir(parents=True)
    docx_path = docx_dir / "brief.docx"

    document = Document()
    document.add_paragraph("项目名称：夏季冰饮新品发布")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "投放城市"
    table.rows[0].cells[1].text = "上海南京西路"
    document.save(docx_path)

    monkeypatch.setattr(document_module.settings, "UPLOAD_DIR", str(upload_dir))
    monkeypatch.setattr(document_module.settings, "OSS_ENABLED", False)
    monkeypatch.setattr(document_module.settings, "AI_API_KEY", "test-key")

    async def _mock_completion(payload, *, timeout=None, attempts=None):
        content = payload["messages"][1]["content"]
        assert "夏季冰饮新品发布" in content
        assert "投放城市 | 上海南京西路" in content
        assert "文档文本" in content
        return {
            "choices": [
                {
                    "message": {
                        "content": (
                            '{"project_name":"夏季冰饮新品发布",'
                            '"city_location":"上海南京西路","remarks":""}'
                        )
                    }
                }
            ]
        }

    monkeypatch.setattr(document_module, "post_chat_completion", _mock_completion)
    result = await document_module.extract_uploaded_brief_documents(
        [
            SimpleNamespace(
                name="brief.docx",
                url="/uploads/site_photos/user-test/brief.docx",
                type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                object_key="",
            )
        ],
        user_id="user-test",
    )

    assert result.updates == {
        "project_name": "夏季冰饮新品发布",
        "city_location": "上海南京西路",
    }
    assert result.filenames == ["brief.docx"]
    assert document_module.BRIEF_DOCUMENT_CONTEXT_MARKER in result.context
    assert "已从文档提取的 Brief 内容" in result.context


@pytest.mark.asyncio
async def test_extract_uploaded_docx_from_oss_preserves_word_extension(monkeypatch):
    from app.services import oss_service

    monkeypatch.setattr(document_module.settings, "OSS_ENABLED", True)
    monkeypatch.setattr(document_module.settings, "AI_API_KEY", "test-key")

    def _mock_download(object_key, target_path):
        assert object_key == "site_photos/user-test/brief.docx"
        assert target_path.endswith(".docx")
        with open(target_path, "wb") as target:
            target.write(b"mock-docx")

    def _mock_parse(path, filename):
        assert path.endswith(".docx")
        assert filename == "brief.docx"
        return [ParsedSection(label="Word正文", page=None, text="项目名称：OSS Word Brief")]

    async def _mock_completion(payload, *, timeout=None, attempts=None):
        assert "OSS Word Brief" in payload["messages"][1]["content"]
        return {
            "choices": [
                {
                    "message": {
                        "content": '{"project_name":"OSS Word Brief","remarks":""}'
                    }
                }
            ]
        }

    monkeypatch.setattr(oss_service, "download_object_to_file", _mock_download)
    monkeypatch.setattr(document_module, "parse_document", _mock_parse)
    monkeypatch.setattr(document_module, "post_chat_completion", _mock_completion)

    result = await document_module.extract_uploaded_brief_documents(
        [
            SimpleNamespace(
                name="brief.docx",
                url="https://example.invalid/brief.docx",
                type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                object_key="site_photos/user-test/brief.docx",
            )
        ],
        user_id="user-test",
    )

    assert result.updates == {"project_name": "OSS Word Brief"}


@pytest.mark.asyncio
async def test_extract_mixed_pdf_and_docx_merges_brief_fields(monkeypatch, tmp_path):
    upload_dir = tmp_path / "uploads"
    document_dir = upload_dir / "site_photos" / "user-test"
    document_dir.mkdir(parents=True)
    (document_dir / "project.pdf").write_bytes(b"%PDF-test")
    (document_dir / "budget.docx").write_bytes(b"mock-docx")

    monkeypatch.setattr(document_module.settings, "UPLOAD_DIR", str(upload_dir))
    monkeypatch.setattr(document_module.settings, "OSS_ENABLED", False)
    monkeypatch.setattr(document_module.settings, "AI_API_KEY", "test-key")

    def _mock_parse(_path, filename):
        text = "项目名称：夏季发布会" if filename.endswith(".pdf") else "预算：30-50万"
        return [ParsedSection(label="正文", page=None, text=text)]

    async def _mock_completion(payload, *, timeout=None, attempts=None):
        content = payload["messages"][1]["content"]
        if "夏季发布会" in content:
            reply = '{"project_name":"夏季发布会"}'
        else:
            reply = '{"budget":"30-50万"}'
        return {"choices": [{"message": {"content": reply}}]}

    monkeypatch.setattr(document_module, "parse_document", _mock_parse)
    monkeypatch.setattr(document_module, "post_chat_completion", _mock_completion)

    result = await document_module.extract_uploaded_brief_documents(
        [
            SimpleNamespace(
                name="project.pdf",
                url="/uploads/site_photos/user-test/project.pdf",
                type="application/pdf",
                object_key="",
            ),
            SimpleNamespace(
                name="budget.docx",
                url="/uploads/site_photos/user-test/budget.docx",
                type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                object_key="",
            ),
        ],
        user_id="user-test",
    )

    assert result.updates == {
        "project_name": "夏季发布会",
        "budget": "30-50万",
    }
    assert result.filenames == ["project.pdf", "budget.docx"]


@pytest.mark.asyncio
async def test_extract_uploaded_legacy_doc_brief_fields(monkeypatch, tmp_path):
    upload_dir = tmp_path / "uploads"
    doc_dir = upload_dir / "site_photos" / "user-test"
    doc_dir.mkdir(parents=True)
    (doc_dir / "legacy-brief.doc").write_bytes(b"mock-legacy-doc")

    monkeypatch.setattr(document_module.settings, "UPLOAD_DIR", str(upload_dir))
    monkeypatch.setattr(document_module.settings, "OSS_ENABLED", False)
    monkeypatch.setattr(document_module.settings, "AI_API_KEY", "test-key")

    def _mock_parse(path, filename):
        assert path.endswith(".doc")
        assert filename == "legacy-brief.doc"
        return [
            ParsedSection(
                label="Word正文",
                page=None,
                text="项目名称：旧版 Word Brief；预算：30-50万",
            )
        ]

    async def _mock_completion(payload, *, timeout=None, attempts=None):
        content = payload["messages"][1]["content"]
        assert "旧版 Word Brief" in content
        return {
            "choices": [
                {
                    "message": {
                        "content": (
                            '{"project_name":"旧版 Word Brief",'
                            '"budget":"30-50万","remarks":""}'
                        )
                    }
                }
            ]
        }

    monkeypatch.setattr(document_module, "parse_document", _mock_parse)
    monkeypatch.setattr(document_module, "post_chat_completion", _mock_completion)
    result = await document_module.extract_uploaded_brief_documents(
        [
            SimpleNamespace(
                name="legacy-brief.doc",
                url="/uploads/site_photos/user-test/legacy-brief.doc",
                type="application/msword",
                object_key="",
            )
        ],
        user_id="user-test",
    )

    assert result.updates == {
        "project_name": "旧版 Word Brief",
        "budget": "30-50万",
    }
    assert result.failures == []


@pytest.mark.asyncio
async def test_legacy_doc_parse_failure_returns_word_message(monkeypatch, tmp_path):
    upload_dir = tmp_path / "uploads"
    doc_dir = upload_dir / "site_photos" / "user-test"
    doc_dir.mkdir(parents=True)
    (doc_dir / "broken.doc").write_bytes(b"broken-legacy-doc")

    monkeypatch.setattr(document_module.settings, "UPLOAD_DIR", str(upload_dir))
    monkeypatch.setattr(document_module.settings, "OSS_ENABLED", False)
    monkeypatch.setattr(document_module.settings, "AI_API_KEY", "test-key")
    monkeypatch.setattr(
        document_module,
        "parse_document",
        lambda *_args: (_ for _ in ()).throw(
            DocumentParseError("旧版 Word 解析失败")
        ),
    )

    result = await document_module.extract_uploaded_brief_documents(
        [
            SimpleNamespace(
                name="broken.doc",
                url="/uploads/site_photos/user-test/broken.doc",
                type="application/msword",
                object_key="",
            )
        ],
        user_id="user-test",
    )

    assert result.failures == ["broken.doc：Word 解析失败"]


@pytest.mark.asyncio
async def test_document_updates_require_confirmation_before_writing_brief_state(monkeypatch, tmp_path):
    from app.services.ai_brief_state import update_agent_state_from_message

    monkeypatch.setattr("app.services.ai_brief_state.settings.LOG_DIR", str(tmp_path))
    monkeypatch.setattr("app.services.ai_brief_state.settings.AI_API_KEY", "test-key")

    async def _mock_confirmation(payload, **_kwargs):
        assert "pending_document_brief" in payload["messages"][1]["content"]
        return {"choices": [{"message": {"content": '{"action":"confirmed","updates":{}}'}}]}

    monkeypatch.setattr("app.services.ai_brief_state.post_chat_completion", _mock_confirmation)

    state = await update_agent_state_from_message(
        session_id="pdf-session",
        user_id="pdf-user",
        business_type="ai_3d_custom",
        message=(
            "[已上传文件: brief.pdf]\n\n"
            "[PDF Brief解析内容]\n"
            "文件：brief.pdf\n"
            "已从 PDF 提取的 Brief 内容：\n"
            "- 投放城市与媒体位置：杭州湖滨银泰"
        ),
        history=[],
        source_message_id="pdf-message",
        memory_hints={},
        document_updates={"city_location": "杭州湖滨银泰"},
    )

    field = state["brief_state"]["fields"]["city_location"]
    assert field["value"] == "杭州湖滨银泰"
    assert state["brief_state"]["fields"]["site_photos"]["value"] == ""
    assert state["pending_document_brief"]["updates"] == {"city_location": "杭州湖滨银泰"}
    context_message = state["agent_context_window"]["messages"][-1]["content"]
    assert "用户上传了文档资料" in context_message
    assert "杭州湖滨银泰" not in context_message

    reviewed = await update_agent_state_from_message(
        session_id="pdf-session",
        user_id="pdf-user",
        business_type="ai_3d_custom",
        message="确认",
        history=[],
        source_message_id="pdf-confirmation",
        memory_hints={},
    )

    reviewed_field = reviewed["brief_state"]["fields"]["city_location"]
    assert reviewed_field["value"] == "杭州湖滨银泰"
    assert reviewed["pending_document_brief"] is None
    assert reviewed["document_brief_confirmation"]["status"] == "reviewed_no_changes"


@pytest.mark.asyncio
async def test_document_revision_updates_candidate_and_waits_for_new_confirmation(monkeypatch, tmp_path):
    from app.services.ai_brief_state import update_agent_state_from_message

    monkeypatch.setattr("app.services.ai_brief_state.settings.LOG_DIR", str(tmp_path))
    monkeypatch.setattr("app.services.ai_brief_state.settings.AI_API_KEY", "test-key")

    async def _mock_completion(payload, **_kwargs):
        input_payload = payload["messages"][1]["content"]
        if "投放地应为上海南京东路" not in input_payload:
            return {"choices": [{"message": {"content": '{"action":"confirmed","updates":{}}'}}]}
        assert "杭州湖滨银泰" in input_payload
        return {
            "choices": [
                {
                    "message": {
                        "content": '{"action":"revised","updates":{"city_location":"上海南京东路"}}'
                    }
                }
            ]
        }

    monkeypatch.setattr("app.services.ai_brief_state.post_chat_completion", _mock_completion)

    await update_agent_state_from_message(
        session_id="revision-session",
        user_id="pdf-user",
        business_type="ai_3d_custom",
        message="[已上传文件: brief.pdf]",
        source_message_id="pdf-message",
        document_updates={"city_location": "杭州湖滨银泰", "project_name": "春季发布会"},
        document_filenames=["brief.pdf"],
    )

    revised = await update_agent_state_from_message(
        session_id="revision-session",
        user_id="pdf-user",
        business_type="ai_3d_custom",
        message="投放地应为上海南京东路",
        source_message_id="revision-message",
    )

    assert revised["brief_state"]["fields"]["city_location"]["value"] == "上海南京东路"
    assert revised["brief_state"]["fields"]["project_name"]["value"] == "春季发布会"
    assert revised["pending_document_brief"] is None
    assert revised["document_brief_confirmation"]["status"] == "revised"
